import os
import sys
import pandas as pd
from tqdm import tqdm
from openai import OpenAI

from docx import Document
import json

# === CONFIGURATION ===
API_KEY = "Here Anouar put the api key"
DATA_DIR = "data"
OUTPUT_DIR = "generated_reports"
os.makedirs(OUTPUT_DIR, exist_ok=True)

client = OpenAI(api_key=API_KEY)

# === Format demandé ===
if len(sys.argv) < 2 or sys.argv[1] not in ["docx", "excel", "json", "md"]:
    print("❌ Veuillez spécifier un format : docx | excel | json | md")
    sys.exit(1)
export_format = sys.argv[1]

# === Lire les fichiers Excel
dataframes = []
for filename in os.listdir(DATA_DIR):
    if filename.endswith(".xlsx"):
        df = pd.read_excel(os.path.join(DATA_DIR, filename))
        dataframes.append(df)
df_all = pd.concat(dataframes, ignore_index=True)

# === Collecte de résultats
results = []

for idx, row in tqdm(df_all.iterrows(), total=df_all.shape[0]):
    interface = str(row.get("Interface", "N/A"))
    title = str(row.get("Test Case Description", "N/A")).strip()
    details = str(row.get("Test Details", "")).strip()
    logs = str(row.get("Logs", "")).strip()

    decision_prompt = f"""
You are a cybersecurity expert reviewing penetration test logs.
Determine if the log indicates a vulnerability or if the component resisted the attack.
Only return one word: VULNERABILITY or SUCCESS.
Log: {logs}
"""
    decision_response = client.chat.completions.create(
        model="gpt-4-turbo",
        messages=[
            {"role": "system", "content": "You are a cybersecurity expert classifying logs."},
            {"role": "user", "content": decision_prompt}
        ]
    )
    decision = decision_response.choices[0].message.content.strip().upper()

    if decision == "SUCCESS":
        results.append({
            "Interface": interface,
            "Test Name": title,
            "Result": "SUCCESS",
            "Message": "The system is resilient. No vulnerability found."
        })
    else:
        prompt = f"""
You are a cybersecurity analyst. Analyze the following penetration test logs and return a detailed vulnerability summary with the following fields:
- CVSS
- Risk level
- Description
- Risks
- Complexity
- Priority
- CWE/CVE reference
- Reference URLs

Interface: {interface}
Test Name: {title}
Details: {details}
Logs: {logs}
"""
        response = client.chat.completions.create(
            model="gpt-4-turbo",
            messages=[
                {"role": "system", "content": "You are a cybersecurity expert writing a professional pentest report."},
                {"role": "user", "content": prompt}
            ]
        )
        summary = response.choices[0].message.content
        record = {"Interface": interface, "Test Name": title, "Result": "VULNERABILITY"}
        for line in summary.splitlines():
            if ":" in line:
                key, val = line.split(":", 1)
                record[key.strip()] = val.strip()
        results.append(record)

# === Export
output_file = os.path.join(OUTPUT_DIR, f"vuln_report.{export_format}")

if export_format == "excel":
    pd.DataFrame(results).to_excel(output_file, index=False)
elif export_format == "json":
    with open(output_file, "w", encoding="utf-8") as f:
        json.dump(results, f, indent=2)
elif export_format == "md":
    with open(output_file, "w", encoding="utf-8") as f:
        for idx, entry in enumerate(results, 1):
            f.write(f"## 3.{idx} {entry['Test Name']}\n")
            for k, v in entry.items():
                if k not in ["Test Name"]:
                    f.write(f"- **{k}**: {v}\n")
            f.write("\n---\n\n")
elif export_format == "docx":
    doc = Document()
    doc.add_heading("📄 VAPT Auto-Generated Vulnerability Report", 0)
    for idx, entry in enumerate(results, 1):
        doc.add_heading(f"3.{idx} {entry['Test Name']}", level=2)
        table = doc.add_table(rows=1, cols=2)
        table.style = "Table Grid"
        table.cell(0, 0).text = "Clé"
        table.cell(0, 1).text = "Valeur"
        for k, v in entry.items():
            if k not in ["Test Name"]:
                row = table.add_row().cells
                row[0].text = k
                row[1].text = str(v)
        doc.add_paragraph()
    doc.save(output_file)

print(f"✅ Rapport généré : {output_file}")
