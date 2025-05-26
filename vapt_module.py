#!/usr/bin/env python3
# coding: utf-8
"""
VAPT Auto-Report Generator (function version)
– ajoute Summary Results & Conclusion –
"""

import os
import sys
import json
import pandas as pd
from tqdm import tqdm
from docx import Document
from llama_cpp import Llama

# === CONFIGURATION ============================================================
DEFAULT_MODEL_PATH = "./openhermes-2.5-mistral-7b.Q5_K_M.gguf"
DATA_DIR   = "data"
OUTPUT_DIR = "generated_reports"
CTX_SIZE   = 2048
os.makedirs(OUTPUT_DIR, exist_ok=True)

# === LLM WRAPPER ==============================================================

def load_model(model_path: str = DEFAULT_MODEL_PATH) -> Llama:
    """Charge le modèle GGUF local avec llama-cpp."""
    return Llama(model_path=model_path, n_ctx=CTX_SIZE)

def ask_llm(llm: Llama, prompt: str, max_tokens: int = 512) -> str:
    """Envoie un prompt au modèle et renvoie la réponse nettoyée."""
    out = llm(f"[INST] {prompt.strip()} [/INST]", max_tokens=max_tokens)
    return out["choices"][0]["text"].strip()

# === GÉNÉRATEUR DE RAPPORT ====================================================

def generate_vapt_report(export_format: str = "docx",
                         model_path: str = DEFAULT_MODEL_PATH) -> str:
    """
    Génère un rapport VAPT complet (Summary, Detailed Findings, Conclusion)
    dans le format demandé : docx | excel | json | md.
    """
    # 1. LLM
    llm = load_model(model_path)

    # 2. Charger toutes les feuilles Excel
    dataframes = []
    for filename in os.listdir(DATA_DIR):
        if filename.endswith(".xlsx"):
            df = pd.read_excel(os.path.join(DATA_DIR, filename))
            dataframes.append(df)

    if not dataframes:
        raise FileNotFoundError("Aucun fichier .xlsx trouvé dans le dossier 'data'.")

    df_all = pd.concat(dataframes, ignore_index=True)

    # 3. Prompts Summary & Conclusion
    summary_prompt = """
You are a senior penetration tester. Based on the following aggregated logs,
write a concise Summary Results section for an executive VAPT report. Mention:
- total number of vulnerabilities and their severities
- overall security level (Low, Moderate, High)
- 3–4 key observations

Aggregated Logs:
""" + df_all["Logs"].dropna().astype(str).str.cat(sep="\n---\n")

    summary_results = ask_llm(llm, summary_prompt, max_tokens=512)

    conclusion_prompt = """
You are a senior penetration tester. Draft the Conclusion section of the report:
- recap the assessment
- highlight the most critical weaknesses
- propose a remediation strategy (quick wins ➜ long term)
- give 3 high-level recommendations
"""
    conclusion_text = ask_llm(llm, conclusion_prompt, max_tokens=512)

    # 4. Analyse détaillée ligne par ligne
    results = []
    for _, row in tqdm(df_all.iterrows(), total=df_all.shape[0]):
        interface = str(row.get("Interface", "N/A")).strip()
        title     = str(row.get("Test Case Description", "N/A")).strip()
        details   = str(row.get("Test Details", "")).strip()
        logs      = str(row.get("Logs", "")).strip()

        decision_prompt = f"""
You are a cybersecurity expert reviewing penetration test logs.
Return only ONE word: VULNERABILITY or SUCCESS.

Log:
{logs}
"""
        decision = ask_llm(llm, decision_prompt, max_tokens=4).upper()

        if "SUCCESS" in decision:
            results.append({
                "Interface": interface,
                "Test Name": title,
                "Result": "SUCCESS",
                "Message": "The system is resilient. No vulnerability found."
            })
        else:
            vuln_prompt = f"""
You are a cybersecurity analyst. Analyse the log and produce a summary with:

- CVSS
- Risk level
- Description
- Risks
- Complexity
- Priority
- CWE/CVE reference
- Reference URLs

Interface: {interface}
Test Name: {title}
Details: {details}
Logs: {logs}
"""
            summary = ask_llm(llm, vuln_prompt, max_tokens=512)

            record = {"Interface": interface,
                      "Test Name": title,
                      "Result": "VULNERABILITY"}

            for line in summary.splitlines():
                if ":" in line:
                    key, val = line.split(":", 1)
                    record[key.strip()] = val.strip()
            results.append(record)

    # 5. Export
    output_file = os.path.join(OUTPUT_DIR, f"vuln_report.{export_format.lower()}")

    if export_format == "excel":
        with pd.ExcelWriter(output_file) as writer:
            pd.DataFrame(results).to_excel(writer,
                                           sheet_name="Detailed Findings",
                                           index=False)
            pd.DataFrame([{"Summary Results": summary_results}]
                         ).to_excel(writer, sheet_name="Summary", index=False)
            pd.DataFrame([{"Conclusion": conclusion_text}]
                         ).to_excel(writer, sheet_name="Conclusion", index=False)

    elif export_format == "json":
        with open(output_file, "w", encoding="utf-8") as f:
            json.dump({
                "summary_results": summary_results,
                "detailed_findings": results,
                "conclusion": conclusion_text
            }, f, indent=2)

    elif export_format == "md":
        with open(output_file, "w", encoding="utf-8") as f:
            f.write("# 📄 VAPT Auto-Generated Vulnerability Report\n\n")
            f.write("## 2. Summary Results\n")
            f.write(summary_results + "\n\n---\n\n")
            f.write("## 3. Detailed Findings\n")
            for idx, entry in enumerate(results, 1):
                f.write(f"### 3.{idx} {entry['Test Name']}\n")
                for k, v in entry.items():
                    if k != "Test Name":
                        f.write(f"- **{k}**: {v}\n")
                f.write("\n---\n\n")
            f.write("## 4. Conclusion\n")
            f.write(conclusion_text + "\n")

    elif export_format == "docx":
        doc = Document()
        doc.add_heading("📄 VAPT Auto-Generated Vulnerability Report", 0)

        # Summary (section 2)
        doc.add_heading("2. Summary Results", level=1)
        doc.add_paragraph(summary_results)

        # Detailed findings (section 3)
        doc.add_heading("3. Detailed Findings", level=1)
        for idx, entry in enumerate(results, 1):
            doc.add_heading(f"3.{idx} {entry['Test Name']}", level=2)
            table = doc.add_table(rows=1, cols=2)
            table.style = "Table Grid"
            table.cell(0, 0).text = "Key"
            table.cell(0, 1).text = "Value"
            for k, v in entry.items():
                if k != "Test Name":
                    row = table.add_row().cells
                    row[0].text = k
                    row[1].text = str(v)
            doc.add_paragraph()

        # Conclusion (section 4)
        doc.add_heading("4. Conclusion", level=1)
        doc.add_paragraph(conclusion_text)

        doc.save(output_file)

    else:
        raise ValueError("Format non supporté : choisir docx | excel | json | md")

    print(f"✅ Rapport généré : {output_file}")
    return output_file

# === EXECUTION DIRECTE ========================================================
if __name__ == "__main__":
    if len(sys.argv) < 2 or sys.argv[1].lower() not in {"docx", "excel", "json", "md"}:
        print("❌ Veuillez spécifier un format : docx | excel | json | md")
        sys.exit(1)
    generate_vapt_report(export_format=sys.argv[1].lower())
